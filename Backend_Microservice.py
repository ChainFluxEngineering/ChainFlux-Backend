from fpdf import FPDF
from flask import Flask, request, jsonify
from flask_cors import CORS
from langchain.text_splitter import RecursiveCharacterTextSplitter
import boto3
import os
import openai
from PyPDF2 import PdfReader
import numpy as np
from docx import Document
from langchain.llms import OpenAI
from langchain.document_loaders import PyPDFLoader
from langchain.chains.summarize import load_summarize_chain
from langchain.embeddings.openai import OpenAIEmbeddings
import pinecone
from langchain.vectorstores import Pinecone
import tempfile
import hashlib
from langchain.document_loaders import DirectoryLoader
from langchain.prompts import PromptTemplate
from langchain.chains.question_answering import load_qa_chain
from langchain.chat_models import ChatOpenAI
from langchain.document_loaders import TextLoader
from datetime import datetime
from flask_pymongo import PyMongo
from langchain.prompts import ChatPromptTemplate
from langchain_community.chat_models import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
import pymongo
from flask import jsonify
import json


os.environ["OPENAI_API_KEY"] = "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"
app = Flask(__name__)
CORS(app)
s3_client = boto3.client('s3')
s3_bucket = 'janak-mvp'
llm = ChatOpenAI(temperature=0)


client = pymongo.MongoClient(
    "mongodb+srv://vyomgoyal2003:vyom2003@janak-ai.qvnw4fy.mongodb.net/?retryWrites=true&w=majority")
db = client.test


# Access the database


def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


def read_word_file(file_path):
    doc = Document(file_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text


def read_pdf_file(file_path):
    textfull = ""
    pdf_reader = PdfReader(file_path)
    for page in pdf_reader.pages:
        textfull += 'New Page\n' + page.extract_text()
    return textfull


@app.route('/upload-pdf', methods=['POST'])
def upload_pdf():
    try:

        # Get the uploaded file from the request
        pdf_file = request.files['pdf']
        s3_folder = request.form['folder']
        user_id = request.form['user_id']

        current_time = datetime.now()
        pdf_id = current_time.strftime('%Y-%m-%d_%H-%M-%S')

        # print(text)

        # Check if the file is empty
        if not pdf_file or not s3_folder:
            return jsonify({'error': 'File and folder name are required'}), 400
        # Create a temporary file to store the content of the PDF
        temp_pdf_file = tempfile.NamedTemporaryFile(delete=False)
        pdf_content = pdf_file.read()
        temp_pdf_file.write(pdf_content)
        temp_pdf_file.close()

        # Specify the S3 folder and file name
        s3_file = pdf_file.filename

        # Upload the temporary file to S3
        s3_client.upload_file(temp_pdf_file.name, s3_bucket,
                              os.path.join(s3_folder, s3_file))

        response_data = {
            'message': f'Document uploaded successfully to s3folder: {s3_folder}'}
        return jsonify(response_data), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/create-embedding', methods=['POST'])
def create_embedding():
    try:
        pdf_file = request.files['pdf']
        user_id = request.form['user_id']
        current_time = datetime.now()
        pdf_id = current_time.strftime('%Y-%m-%d_%H-%M-%S')
        print(pdf_file)
        textfull = ""

        # Determine the file type based on the extension
        _, extension = os.path.splitext(pdf_file.filename)

        if extension.lower() == '.pdf':
            # PDF file
            pdf_reader = PdfReader(pdf_file)
            for page in pdf_reader.pages:
                textfull += 'New Page\n' + page.extract_text()
        elif extension.lower() == '.docx':
            # Word document (docx)
            doc = Document(pdf_file)
            for paragraph in doc.paragraphs:
                textfull += paragraph.text + '\n'
        elif extension.lower() == '.txt':
            # Text file (txt)
            textfull = pdf_file.read().decode('utf-8')

        # prompt_template = """Write a summary of the following keeping the context and important details intact:

        # {text}"""
        # summary = []
        # text_pages = textfull.split('New Page\n')
        # for page in text_pages:
        #     with open('temp.txt', 'w') as file:
        #         file.write(page)
        #     loader = DirectoryLoader(
        #         './', glob="**/temp.txt", loader_cls=TextLoader)

        #     docs = loader.load_and_split()
        #     prompt = PromptTemplate(
        #         template=prompt_template, input_variables=["text"])
        #     chain = load_summarize_chain(
        #         llm, chain_type="map_reduce", map_prompt=prompt, combine_prompt=prompt, verbose=False)
        #     for i in docs:
        #         summary.append(chain.run([i]))
        # text = '\n'.join(summary)

        # print(text)

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, chunk_overlap=200, length_function=len)
        # docs = text_splitter.split_text(text)
        docsfull = text_splitter.split_text(textfull)
        # print(docs)

        # # Store the embeddings to Pinecone
        pinecone.init(
            api_key='ba7c38c0-27bb-4559-aa3e-627f57096990',
            environment='eu-west4-gcp'
        )
        # index_name = "summary"
        index_name = user_id
        # index_exists = False

        index_exists_full = False

        indexes = pinecone.list_indexes()

        # if index_name in indexes:
        #     index_exists = True

        if index_name in indexes:
            index_exists_full = True

        # if not index_exists:
        #     # Create the index if it doesn't exist
        #     pinecone.create_index(index_name, dimension=1536, metric='cosine')

        if not index_exists_full:
            # Create the index if it doesn't exist
            pinecone.create_index(
                index_name, dimension=1536, metric='cosine')

        # index = pinecone.Index(index_name)

        index_full = pinecone.Index(index_name)

        embeddings = OpenAIEmbeddings(model_name="ada")

        # documents = []

        # for chunk_id, chunk in enumerate(docs):
        #     # Generate embeddings for the text chunk
        #     vec = embeddings.embed_query(chunk)
        #     metadata = {
        #         "pdf_id": pdf_id,
        #         "user_id": user_id,
        #         "text": chunk,
        #     }
        #     document = {
        #         "id": f"{pdf_id}_{user_id}_{chunk_id}",
        #         # Unique ID for each chunk
        #         "values": vec,  # Use the embeddings generated from your text chunk
        #         "metadata": metadata
        #     }
        #     documents.append(document)

        documentsfull = []

        for chunk_id, chunk in enumerate(docsfull):
            # Generate embeddings for the text chunk
            vec = embeddings.embed_query(chunk)
            metadata = {
                "pdf_id": pdf_id,
                "user_id": user_id,
                "text": chunk,
            }
            document = {
                "id": f"{pdf_id}_{user_id}_{chunk_id}",
                # Unique ID for each chunk
                "values": vec,  # Use the embeddings generated from your text chunk
                "metadata": metadata
            }
            documentsfull.append(document)

        # index.upsert(documents)
        index_full.upsert(documentsfull)

        print(pdf_id)

        response_data = {
            'message': f'Index: {index_name} created and embeddings stored succesfully.'}
        return jsonify(response_data), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/query', methods=['POST'])
def query():
    try:
        query = request.form['query']
        user_id = request.form['user_id']

        text = "Hello this is test insert"
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, chunk_overlap=200, length_function=len)
        docs = text_splitter.split_text(text)

        pinecone.init(
            api_key='ba7c38c0-27bb-4559-aa3e-627f57096990',
            environment='eu-west4-gcp'
        )
        # index_name = 'summary'
        index_name_full = user_id
        # index_exists = False
        index_exists_full = False

        indexes = pinecone.list_indexes()

        # if index_name in indexes:
        #     index_exists = True

        if index_name_full in indexes:
            index_exists_full = True

        if not index_exists_full:
            response_data = {'message': 'query unsuccessfully',
                             'answer': "Please upload a File!"}
            return jsonify(response_data), 200

        embeddings = OpenAIEmbeddings(model_name="ada")
        embedding = embeddings.embed_query(docs[0])
        # index = pinecone.Index(index_name)
        # index = Pinecone.from_texts(docs, embeddings, index_name=index_name)

        index_full = pinecone.Index(index_name_full)
        index_full = Pinecone.from_texts(
            docs, embeddings, index_name=index_name_full)

        k = 3
        # Define the desired pdf_id you want to filter by
        # desired_pdf_id = "pdf_6"

        # filter_condition = f"metadata.pdf_id:{desired_pdf_id}"

        # score = False
        # if score:
        #     similar_docs = index.similarity_search_with_score(query, k=k)
        # else:
        #     similar_docs = index.similarity_search(query, k=k)
        # similar_docs

        # metadata_list = []
        # # metadata_list[0]

        # for doc in similar_docs:
        #     metadata = doc.metadata
        #     metadata_list.append(metadata)

        # print(metadata_list)

        # pdf_ids = [item['pdf_id'] for item in metadata_list if item]

        # print(pdf_ids)

        # filter1 = {
        #      "user_id": {"$eq": user_id}
        # }

        # print(filter1)

        score = False
        if score:
            similar_docs_full = index_full.similarity_search_with_score(
                query, k=k)
        else:
            similar_docs_full = index_full.similarity_search(
                query, k=k)

        # metadata_list_full = []

        # for doc in similar_docs_full:
        #     metadata = doc.metadata
        #     metadata_list_full.append(metadata)

        # print(metadata_list_full)

        # model_name = "text-davinci-003"
        model_name = "gpt-3.5-turbo"
        # model_name = "gpt-4"
        llm = OpenAI(model_name=model_name)

        chain = load_qa_chain(llm, chain_type="stuff")

        answer = chain.run(input_documents=similar_docs_full, question=query)

        response_data = {'message': 'query successfully', 'answer': answer}

        return jsonify(response_data), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/delete', methods=['POST'])
def delete_user():
    try:

        user_id = request.form['user_id']

        pinecone.init(
            api_key='ba7c38c0-27bb-4559-aa3e-627f57096990',
            environment='eu-west4-gcp'
        )
        # index_name = 'summary'
        index_name_full = user_id
        index_exists = False
        index_exists_full = False

        indexes = pinecone.list_indexes()

        # if index_name in indexes:
        #     index_exists = True

        if index_name_full in indexes:
            index_exists_full = True

        # index = pinecone.Index(index_name)

        index_full = pinecone.Index(index_name_full)

        # filter_condition = f"metadata.pdf_id:{desired_pdf_id}"

        # filter1 = {
        #      "user_id": {"$eq": user_id}
        # }

        # print(filter1)


#         index.delete(
#     filter= filter1


# )
        # index_full.delete(
        #     filter = filter1
        # )

        # metadata_list_full = []

        # for doc in similar_docs_full:
        #     metadata = doc.metadata
        #     metadata_list_full.append(metadata)

        # print(metadata_list_full)

        # model_name = "text-davinci-003"
        pinecone.delete_index(index_name_full)

        response_data = {'message': 'query successfully',
                         'answer': f'User {user_id} Deleted Succesfully'}

        return jsonify(response_data), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/listfile', methods=['POST'])
def list_file():
    try:

        s3_folder = request.form['folder']

        response = s3_client.list_objects_v2(
            Bucket=s3_bucket, Prefix=s3_folder)

        object_keys = [obj['Key'] for obj in response.get('Contents', [])]

        num_files = len(object_keys)

        response_data = {
            'message': 'Query successful',
            's3_folder': s3_folder,
            'object_keys': object_keys,
            'no_files': num_files
        }

        return jsonify(response_data), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500



@app.route('/get-names', methods=['POST'])
def get_names():
    try:
        # Access the 'agree_metadata' collection
        collection = db.agree_metadata  # Use your collection name

        # Retrieve all documents from the collection
        names = collection.find({}, {'name': 1})

        # Convert the cursor to a list of names
        name_list = [doc['name'] for doc in names]

        # response_data = {'message': 'Query successful', 'names': name_list}

        user_query = request.form['user_query']

        # Use OpenAI to generate a response based on the user's query
        prompt = f"Find the most relevant agreement type for the following query: {user_query} from the list of agreements {name_list}. Give the name exactly as in array and give the answer in one word."
        print(prompt)
        openai.api_key = "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"
        completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ]
        )
        type_agreement = completion['choices'][0]['message']['content']



        selected_agreement = collection.find_one({'name': type_agreement})

        if not selected_agreement:
            return jsonify({'error': f'Agreement type "{type_agreement}" not found'}), 404

        # Extract data associated with the selected agreement type
        data = selected_agreement.get('data', {})

        # Convert data to a string
        data = json.loads(data)
        sorted_data = sorted(data.items(), key=lambda x: [int(num) for num in x[0].split('.')])

        # Create a prompt using the data of the selected agreement type
        generated_agreement = ""
        for clause_number, clause_data in sorted_data:
            # Extract key phrases and purpose for the current clause
            print(clause_number, "\n")
            key_phrases = clause_data.get('key_phrases', [])
            purpose = clause_data.get('purpose', '')

            # Construct a prompt for the current clause
            prompt = f"Generate the agreement clause for the following key phrases: {', '.join(key_phrases)}. Purpose: {purpose}"

            # Use OpenAI to generate the clause for the current key phrases
            completion = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": prompt}
                ]
            )
            generated_clause = completion['choices'][0]['message']['content']

            # Combine the generated clause with the overall agreement
            generated_agreement += f"\n\nClause {clause_number}\n{generated_clause}\n"
        print(generated_agreement)
        response_data = {'message': 'Query successful',
                         'generated_agreement': generated_agreement}

        return jsonify(response_data), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
